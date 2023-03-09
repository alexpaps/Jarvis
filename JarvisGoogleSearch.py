def JarvisWebSearch(query):
    from googlesearch import search

    #query = "ronaldos age"
    num_results = 0
    for result in search(query):
       # print(result)
        num_results += 1
        if num_results == 1:
            break
    return result

def JarvisInfSearch(url,SearchName):
    import requests
    from bs4 import BeautifulSoup
    from docx import Document

    # Retrieve website content
    #url = 'https://en.wikipedia.org/wiki/Cristiano_Ronaldo'
    response = requests.get(url)
    html_content = response.text

    # Extract text from HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    text = soup.get_text()

    # Save text to Word document
    doc = Document()
    doc.add_paragraph(text)
    doc.save('{}.docx'.format(SearchName))



def Jarvisdoc(SearchName):
    from docx import Document
    # Load the Word document
    doc = Document('{}.docx'.format(SearchName))

    lista = []
    TelLista = []

    # Iterate over each paragraph in the document
    for para in doc.paragraphs:
        # Iterate over each line in the paragraph
        for line in para.text.split('\n'):
            # Do something with the line
            lista = line.split()
            if len(lista) > 9:
                TelLista.append(line)

    Maxi = 0
    #Sasw to dekdi orio an exei ligotera apo 8
    try:
        for i in range(4,8):
            if TelLista[i].find('EnglishSlovenčina') != -1:
                continue
            if len(TelLista[i]) > 300:
                Maxi = i
                break
    #Sazw to expect otan den exw plirofories
    except:
        Maxi = 0

    return TelLista[Maxi]

def Translation():
    from googletrans import Translator

    translator = Translator()
    text = "Hello world!"
    translated_text = translator.translate(text, dest='es').text
    print(translated_text)
